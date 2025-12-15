"""
Export modülü - PDF, DOCX, TXT, Markdown export
"""
from datetime import datetime
from pathlib import Path
from typing import Dict, List, Optional
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import markdown


class ExportManager:
    """Export yönetim sınıfı"""
    
    def __init__(self, export_dir: str = "data/exports"):
        self.export_dir = Path(export_dir)
        self.export_dir.mkdir(parents=True, exist_ok=True)
    
    def export_to_txt(self, entry: Dict, filename: Optional[str] = None) -> str:
        """TXT formatında export"""
        if not filename:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"research_{timestamp}.txt"
        
        filepath = self.export_dir / filename
        
        content = self._format_entry(entry, format_type="txt")
        
        with open(filepath, 'w', encoding='utf-8') as f:
            f.write(content)
        
        return str(filepath)
    
    def export_to_markdown(self, entry: Dict, filename: Optional[str] = None) -> str:
        """Markdown formatında export"""
        if not filename:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"research_{timestamp}.md"
        
        filepath = self.export_dir / filename
        
        content = self._format_entry(entry, format_type="markdown")
        
        with open(filepath, 'w', encoding='utf-8') as f:
            f.write(content)
        
        return str(filepath)
    
    def export_to_docx(self, entry: Dict, filename: Optional[str] = None) -> str:
        """DOCX formatında export"""
        if not filename:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"research_{timestamp}.docx"
        
        filepath = self.export_dir / filename
        
        doc = Document()
        
        # Başlık
        title = doc.add_heading('Araştırma Raporu', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Tarih
        date_para = doc.add_paragraph(f"Tarih: {entry.get('timestamp', 'Bilinmiyor')}")
        date_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        # Model bilgisi
        doc.add_heading('Model Bilgisi', level=1)
        doc.add_paragraph(f"Model: {entry.get('model', 'Bilinmiyor')}")
        
        # Prompt
        doc.add_heading('Soru/Prompt', level=1)
        doc.add_paragraph(entry.get('prompt', ''))
        
        # Dosyalar
        files = entry.get('files', [])
        if files:
            doc.add_heading('Eklenen Dosyalar', level=1)
            for file_path in files:
                doc.add_paragraph(f"• {file_path}", style='List Bullet')
        
        # Web arama sonuçları
        web_results = entry.get('web_search_results', [])
        if web_results:
            doc.add_heading('Web Arama Sonuçları', level=1)
            for i, result in enumerate(web_results, 1):
                doc.add_paragraph(f"{i}. {result.get('title', '')}", style='Heading 3')
                doc.add_paragraph(f"URL: {result.get('url', '')}")
                doc.add_paragraph(result.get('snippet', ''))
                doc.add_paragraph()
        
        # Yanıt
        doc.add_heading('Yanıt', level=1)
        doc.add_paragraph(entry.get('response', ''))
        
        # Kaynaklar
        if web_results:
            doc.add_heading('Kaynaklar', level=1)
            for result in web_results:
                doc.add_paragraph(result.get('url', ''), style='List Bullet')
        
        doc.save(filepath)
        return str(filepath)
    
    def export_to_pdf(self, entry: Dict, filename: Optional[str] = None) -> str:
        """PDF formatında export (Markdown'dan dönüştür)"""
        # Markdown export yap, sonra PDF'e çevir (basit yaklaşım)
        # Not: Tam PDF export için reportlab veya weasyprint gerekir
        # Şimdilik Markdown export yapıyoruz
        md_file = self.export_to_markdown(entry, filename.replace('.pdf', '.md') if filename else None)
        return md_file
    
    def _format_entry(self, entry: Dict, format_type: str = "txt") -> str:
        """Entry'yi formatla"""
        lines = []
        
        if format_type == "markdown":
            lines.append("# Araştırma Raporu\n")
            lines.append(f"**Tarih:** {entry.get('timestamp', 'Bilinmiyor')}\n")
            lines.append(f"**Model:** {entry.get('model', 'Bilinmiyor')}\n\n")
            lines.append("## Soru/Prompt\n")
            lines.append(f"{entry.get('prompt', '')}\n\n")
            
            files = entry.get('files', [])
            if files:
                lines.append("## Eklenen Dosyalar\n")
                for file_path in files:
                    lines.append(f"- {file_path}\n")
                lines.append("\n")
            
            web_results = entry.get('web_search_results', [])
            if web_results:
                lines.append("## Web Arama Sonuçları\n")
                for i, result in enumerate(web_results, 1):
                    lines.append(f"### {i}. {result.get('title', '')}\n")
                    lines.append(f"**URL:** {result.get('url', '')}\n")
                    lines.append(f"{result.get('snippet', '')}\n\n")
            
            lines.append("## Yanıt\n")
            lines.append(f"{entry.get('response', '')}\n\n")
            
            if web_results:
                lines.append("## Kaynaklar\n")
                for result in web_results:
                    lines.append(f"- {result.get('url', '')}\n")
        
        else:  # txt
            lines.append("=" * 60)
            lines.append("ARAŞTIRMA RAPORU")
            lines.append("=" * 60)
            lines.append(f"Tarih: {entry.get('timestamp', 'Bilinmiyor')}")
            lines.append(f"Model: {entry.get('model', 'Bilinmiyor')}")
            lines.append("=" * 60)
            lines.append("\nSORU/PROMPT:")
            lines.append("-" * 60)
            lines.append(entry.get('prompt', ''))
            lines.append("\n")
            
            files = entry.get('files', [])
            if files:
                lines.append("EKLENEN DOSYALAR:")
                lines.append("-" * 60)
                for file_path in files:
                    lines.append(f"  • {file_path}")
                lines.append("\n")
            
            web_results = entry.get('web_search_results', [])
            if web_results:
                lines.append("WEB ARAMA SONUÇLARI:")
                lines.append("-" * 60)
                for i, result in enumerate(web_results, 1):
                    lines.append(f"\n{i}. {result.get('title', '')}")
                    lines.append(f"   URL: {result.get('url', '')}")
                    lines.append(f"   {result.get('snippet', '')}")
                lines.append("\n")
            
            lines.append("YANIT:")
            lines.append("-" * 60)
            lines.append(entry.get('response', ''))
            lines.append("\n")
            
            if web_results:
                lines.append("KAYNAKLAR:")
                lines.append("-" * 60)
                for result in web_results:
                    lines.append(f"  • {result.get('url', '')}")
        
        return "\n".join(lines)
    
    def export_multiple(self, entries: List[Dict], format_type: str = "txt", filename: Optional[str] = None) -> str:
        """Birden fazla entry'yi export et"""
        if not filename:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"research_batch_{timestamp}.{format_type if format_type != 'docx' else 'docx'}"
        
        if format_type == "docx":
            return self._export_multiple_docx(entries, filename)
        else:
            content_parts = []
            for entry in entries:
                content_parts.append(self._format_entry(entry, format_type))
                content_parts.append("\n" + "=" * 80 + "\n")
            
            filepath = self.export_dir / filename
            with open(filepath, 'w', encoding='utf-8') as f:
                f.write("\n".join(content_parts))
            
            return str(filepath)
    
    def _export_multiple_docx(self, entries: List[Dict], filename: str) -> str:
        """Birden fazla entry'yi DOCX olarak export et"""
        filepath = self.export_dir / filename
        doc = Document()
        
        doc.add_heading('Toplu Araştırma Raporu', 0)
        
        for i, entry in enumerate(entries, 1):
            doc.add_page_break() if i > 1 else None
            doc.add_heading(f'Araştırma {i}', level=1)
            
            doc.add_paragraph(f"Tarih: {entry.get('timestamp', 'Bilinmiyor')}")
            doc.add_paragraph(f"Model: {entry.get('model', 'Bilinmiyor')}")
            doc.add_heading('Soru/Prompt', level=2)
            doc.add_paragraph(entry.get('prompt', ''))
            doc.add_heading('Yanıt', level=2)
            doc.add_paragraph(entry.get('response', ''))
        
        doc.save(filepath)
        return str(filepath)

